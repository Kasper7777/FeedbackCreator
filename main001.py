import os
from datetime import datetime
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Pt, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Configuration settings
config = {
    'percentage_mark_label': 'Percentage Mark:',
    'tutor_name': 'Dr. John Doe',  # Add the tutor's name
    'module_title': 'Introduction to Programming',  # Add module title
    'module_code': 'CSC101',  # Add module code
    'assignment_title': 'Assignment 1',  # Add assignment title
    'percent_of_module': '25%',  # Add percentage of module
    'font_sizes': {
        'title': 20,
        'year': 12,
        'comment': 14,
        'footer': 9
    }
}


def find_excel_file():
    """Find the first Excel file in the current directory."""
    for file in os.listdir('.'):
        if file.endswith('.xlsx'):
            return file
    return None


def set_cell_border(cell, color='000000', **kwargs):
    """Set cell border for table."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for edge in ('top', 'left', 'bottom', 'right'):
        if edge in kwargs:
            edge_pr = OxmlElement(f'w:{edge}')
            edge_pr.set(qn('w:val'), 'single')
            edge_pr.set(qn('w:sz'), str(kwargs[edge] * 8))  # 8 represents 1pt size
            edge_pr.set(qn('w:space'), '0')
            edge_pr.set(qn('w:color'), color)
            tcPr.append(edge_pr)


def add_centered_heading(doc, text, font_size=18):
    """Add a centered heading."""
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt(font_size)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def create_feedback_form(student_name, student_mark, feedback):
    document = Document()

    # Optionally set document margins
    sections = document.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Title
    add_centered_heading(document, 'ASSIGNMENT FEEDBACK FORM', font_size=config['font_sizes']['title'])
    add_centered_heading(document, '2024-25', font_size=config['font_sizes']['year'])

    # Add Student Info table and set widths
    table = document.add_table(rows=3, cols=4)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    widths = [Inches(1.5), Inches(1.5), Inches(1.5), Inches(1.5)]

    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width

    # Set student information cells
    table.cell(0, 0).text = 'STUDENT:'
    table.cell(0, 1).text = student_name
    table.cell(0, 2).text = 'TUTOR:'
    table.cell(0, 3).text = config['tutor_name']
    table.cell(1, 0).text = 'MODULE TITLE:'
    table.cell(1, 1).text = config['module_title']
    table.cell(1, 2).text = 'MODULE CODE:'
    table.cell(1, 3).text = config['module_code']
    table.cell(2, 0).text = 'ASSIGNMENT:'
    table.cell(2, 1).text = config['assignment_title']
    table.cell(2, 2).text = '% of module:'
    table.cell(2, 3).text = config['percent_of_module']

    # Set table borders
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell, top=1, left=1, bottom=1, right=1)

    document.add_paragraph()  # Add a break

    # Add 'Overall Comment' box and set widths
    add_centered_heading(document, 'OVERALL COMMENT', font_size=config['font_sizes']['comment'])
    comment_table = document.add_table(rows=1, cols=1)
    comment_table.autofit = False
    comment_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    comment_table.cell(0, 0).width = Inches(6)
    comment_table.cell(0, 0).text = feedback if feedback else ' '
    set_cell_border(comment_table.cell(0, 0), top=1, left=1, bottom=1, right=1)
    document.add_paragraph()  # Add a break

    # Marks and Date table and set widths
    table = document.add_table(rows=1, cols=3)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.cell(0, 0).width = Inches(2)
    table.cell(0, 1).width = Inches(2)
    table.cell(0, 2).width = Inches(2)

    table.cell(0, 0).text = config['percentage_mark_label']
    mark_paragraph = table.cell(0, 1).paragraphs[0]
    mark_run = mark_paragraph.add_run(str(student_mark))
    mark_run.bold = True
    table.cell(0, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.cell(0, 2).text = 'Date:'
    # Set today's date
    table.cell(0, 2).paragraphs[0].add_run(datetime.now().strftime('%Y-%m-%d')).bold = True
    # Center alignment for the `cw_label: mark`
    table.cell(0, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.cell(0, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Set borders for 'Marks and Date' table
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell, top=1, left=1, bottom=1, right=1)

    document.add_paragraph()  # Add a break

    # Footer message
    footer = document.add_paragraph()
    footer_run = footer.add_run(
        "NB All marks are provisional until confirmed by a formally constituted Board of Examiners")
    footer_run.font.size = Pt(config['font_sizes']['footer'])  # Smaller font for the footer message
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Save the document
    document.save(f'Assignment_Feedback_Form_{student_name}.docx')


# Find the Excel file in the current directory
excel_file = find_excel_file()
if not excel_file:
    print("No Excel file found in the current directory.")
else:
    # Read student names, marks, and feedback from Excel file
    df = pd.read_excel(excel_file)

    # Assuming the student names are in the first column, marks in the second, and feedback in the third column
    student_data = df.iloc[:, [0, 1, 2]]

    # Create a feedback form for each student
    for index, row in student_data.iterrows():
        student_name = row.iloc[0]
        student_mark = row.iloc[1]
        feedback = row.iloc[2]
        create_feedback_form(student_name, student_mark, feedback)