#enjoy the code
import os
from datetime import datetime
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Pt, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import tkinter as tk
from tkinter import filedialog, messagebox
from tkinter import ttk

# Configuration settings
config = {
    'percentage_mark_label': 'Percentage Mark:',
    'tutor_name': 'Dr. Kazber',
    'module_title': 'Introduction to Game Design',
    'module_code': 'GD101',
    'assignment_title': 'Assignment 1',
    'percent_of_module': '100%',
    'font_sizes': {
        'title': 20,
        'year': 12,
        'comment': 14,
        'footer': 9
    }
}

excel_file_path = None
word_file_path = None


# Updated function to set cell border using proper namespaced attributes
def set_cell_border(cell, **kwargs):
    """
        Sets the border of a table cell in a word document.

        Parameters:
        cell: The table cell to apply borders to.
        kwargs: Keyword arguments specifying border width for 'top', 'left', 'bottom', and 'right' edges.

        Usage:
        Call this function with the table cell and specify the border widths in points.
        For example, set_cell_border(cell, top=2, left=2, bottom=2, right=2).
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for edge in ('top', 'left', 'bottom', 'right'):
        edge_data = kwargs.get(edge, None)
        if edge_data:
            edge_element = OxmlElement(f'w:{edge}')
            edge_element.set(qn('w:val'), 'single')
            edge_element.set(qn('w:sz'), str(edge_data * 8))  # width in eighths of a point
            edge_element.set(qn('w:space'), '0')
            edge_element.set(qn('w:color'), '000000')
            tcPr.append(edge_element)


# Helper function to add a centered heading
def add_centered_heading(doc, text, font_size=18):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt(font_size)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


# Helper function to extract the last table from a .docx document
def extract_last_table_from_docx(doc_path):
    doc = Document(doc_path)
    if doc.tables:
        return doc.tables[-1]
    return None


def append_table_to_document(target_doc, table):
    """Append the given table to the target document without modifications."""
    new_table = target_doc.add_table(rows=len(table.rows), cols=len(table.columns))
    new_table.autofit = False
    for row_idx, row in enumerate(table.rows):
        for cell_idx, cell in enumerate(row.cells):
            new_cell = new_table.cell(row_idx, cell_idx)
            new_cell.text = cell.text
            set_cell_border(new_cell, top=1, left=1, bottom=1, right=1)
            # Copying formatting
            for paragraph in cell.paragraphs:
                new_p = new_cell.paragraphs[0]
                for run in paragraph.runs:
                    new_run = new_p.add_run(run.text)
                    if run.bold:
                        new_run.bold = True
                    if run.italic:
                        new_run.italic = True
                    new_run.font.name = run.font.name
                    new_run.font.size = run.font.size
                new_p.alignment = paragraph.alignment


def create_feedback_form(student_name, student_mark, feedback, extra_table):
    doc = Document()

    # Set author to tutor's name
    core_properties = doc.core_properties
    core_properties.author = config['tutor_name']

    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Title
    add_centered_heading(doc, 'ASSIGNMENT FEEDBACK FORM', font_size=config['font_sizes']['title'])
    add_centered_heading(doc, '2024-25', font_size=config['font_sizes']['year'])

    # Add Student Info table
    table = doc.add_table(rows=3, cols=4)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in table.rows:
        for cell in row.cells:
            cell.width = Inches(1.5)

    # Set student information cells
    table.cell(0, 0).text = 'STUDENT:'
    table.cell(0, 1).text = student_name
    table.cell(0, 2).text = 'TUTOR:'
    table.cell(0, 3).text = config['tutor_name']
    table.cell(1, 0).text = 'MODULE TITLE:'
    table.cell(1, 1).text = config['module_title']
    table.cell(1, 2).text = 'MODULE CODE:'
    table.cell(1, 3).text = config['module_code']
    table.cell(2, 0).text = 'ASSIGNMENT:'
    table.cell(2, 1).text = config['assignment_title']
    table.cell(2, 2).text = '% of module:'
    table.cell(2, 3).text = config['percent_of_module']

    # Set table borders
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell, top=1, left=1, bottom=1, right=1)

    doc.add_paragraph()  # Add a break

    # Add "Overall Comment" box
    add_centered_heading(doc, 'OVERALL COMMENT', font_size=config['font_sizes']['comment'])
    comment_table = doc.add_table(rows=1, cols=1)
    comment_table.autofit = False
    comment_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    comment_table.cell(0, 0).width = Inches(6)
    comment_table.cell(0, 0).text = feedback if feedback else ' '
    set_cell_border(comment_table.cell(0, 0), top=1, left=1, bottom=1, right=1)
    doc.add_paragraph()  # Add a break

    # Add Marks and Date table
    table = doc.add_table(rows=1, cols=3)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for cell in table.rows[0].cells:
        cell.width = Inches(2)

    table.cell(0, 0).text = config['percentage_mark_label']
    mark_paragraph = table.cell(0, 1).paragraphs[0]
    mark_run = mark_paragraph.add_run(str(student_mark))
    mark_run.bold = True
    table.cell(0, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.cell(0, 2).text = 'Date:'
    table.cell(0, 2).paragraphs[0].add_run(datetime.now().strftime('%Y-%m-%d')).bold = True

    # Set table borders for Marks and Date table
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell, top=1, left=1, bottom=1, right=1)

    doc.add_paragraph()  # Add a break

    # Footer message
    footer = doc.add_paragraph()
    footer_run = footer.add_run(
        "NB All marks are provisional until confirmed by a formally constituted Board of Examiners")
    footer_run.font.size = Pt(config['font_sizes']['footer'])
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()  # Add a break

    # Add a page break after the footer message
    doc.add_page_break()

    # Append the extra table if it exists
    if extra_table:
        append_table_to_document(doc, extra_table)

    # Save the document
    file_name = f'Assignment_Feedback_Form_{student_name}.docx'
    doc.save(file_name)
    print(f"Saved: {file_name}")


def process_files():
    global excel_file_path, word_file_path
    if not excel_file_path:
        messagebox.showerror("Error", "Please select an Excel file.")
        return

    extra_table = None
    if word_checkbox_var.get() and word_file_path:
        extra_table = extract_last_table_from_docx(word_file_path)

    # Read student names, marks, and feedback from the Excel file
    try:
        df = pd.read_excel(excel_file_path)
        student_data = df.iloc[:, [0, 1, 2]]

        # Create a feedback form for each student
        for index, row in student_data.iterrows():
            student_name = row.iloc[0]
            student_mark = row.iloc[1]
            feedback = row.iloc[2]
            create_feedback_form(student_name, student_mark, feedback, extra_table)
    except Exception as e:
        messagebox.showerror("Error", f"Failed to process files: {e}")


def select_excel_file():
    global excel_file_path
    file_path = filedialog.askopenfilename(filetypes=[("Excel files", "*.xlsx")])
    if file_path:
        excel_file_path = file_path
        excel_label.config(text=f"Excel file: {os.path.basename(file_path)}")


def select_word_file():
    global word_file_path
    file_path = filedialog.askopenfilename(filetypes=[("Word files", "*.docx")])
    if file_path:
        word_file_path = file_path
        word_label.config(text=f"Word file: {os.path.basename(file_path)}")


def update_config():
    config['tutor_name'] = tutor_name_var.get()
    config['module_title'] = module_title_var.get()
    config['module_code'] = module_code_var.get()
    config['assignment_title'] = assignment_title_var.get()
    config['percent_of_module'] = percent_of_module_var.get()
    messagebox.showinfo("Success", "Configuration updated successfully.")


def toggle_word_file_selection():
    if word_checkbox_var.get():
        word_button.config(state=tk.NORMAL)
        word_label.config(state=tk.NORMAL)
    else:
        word_button.config(state=tk.DISABLED)
        word_label.config(state=tk.DISABLED)


def open_website(event):
    import webbrowser
    webbrowser.open_new("https://kasper7777.github.io/")


# GUI setup
root = tk.Tk()
root.title("Assignment Feedback Form Generator")

# Configuration frame
config_frame = tk.LabelFrame(root, text="Configuration", padx=10, pady=10)
config_frame.pack(padx=10, pady=10, fill="x")

tutor_name_var = tk.StringVar(value=config['tutor_name'])
module_title_var = tk.StringVar(value=config['module_title'])
module_code_var = tk.StringVar(value=config['module_code'])
assignment_title_var = tk.StringVar(value=config['assignment_title'])
percent_of_module_var = tk.StringVar(value=config['percent_of_module'])

tk.Label(config_frame, text="Tutor Name:").grid(row=0, column=0)
tk.Entry(config_frame, textvariable=tutor_name_var).grid(row=0, column=1)

tk.Label(config_frame, text="Module Title:").grid(row=1, column=0)
tk.Entry(config_frame, textvariable=module_title_var).grid(row=1, column=1)

tk.Label(config_frame, text="Module Code:").grid(row=2, column=0)
tk.Entry(config_frame, textvariable=module_code_var).grid(row=2, column=1)

tk.Label(config_frame, text="Assignment Title:").grid(row=3, column=0)
tk.Entry(config_frame, textvariable=assignment_title_var).grid(row=3, column=1)

tk.Label(config_frame, text="Percent of Module:").grid(row=4, column=0)
tk.Entry(config_frame, textvariable=percent_of_module_var).grid(row=4, column=1)

tk.Button(config_frame, text="Update Configuration", command=update_config).grid(row=5, columnspan=2, pady=10)

# File selection frame
file_frame = tk.LabelFrame(root, text="Select Files", padx=10, pady=10)
file_frame.pack(padx=10, pady=10, fill="x")

excel_button = tk.Button(file_frame, text="Select Excel File", command=select_excel_file)
excel_button.pack(pady=5)
excel_label = tk.Label(file_frame, text="No Excel file selected")
excel_label.pack(pady=5)

# Optional Word file selection
word_checkbox_var = tk.BooleanVar(value=False)
word_checkbox = tk.Checkbutton(file_frame, text="Include Rubric (Word Doc)", variable=word_checkbox_var,
                               command=toggle_word_file_selection)
word_checkbox.pack(pady=5)

word_button = tk.Button(file_frame, text="Select Word File", command=select_word_file, state=tk.DISABLED)
word_button.pack(pady=5)
word_label = tk.Label(file_frame, text="No Word file selected", state=tk.DISABLED)
word_label.pack(pady=5)

# Process button
process_button = tk.Button(root, text="Process Files", command=process_files)
process_button.pack(pady=10)

# Link and Copyright
footer_frame = tk.Frame(root)
footer_frame.pack(pady=10)

link = ttk.Label(footer_frame, text="Visit our website: kasper7777.github.io", foreground="blue", cursor="hand2")
link.pack(side="top")
link.bind("<Button-1>", open_website)

copyright_label = tk.Label(footer_frame, text="©2025 Kestrel Kinetics Research & Technology. All Rights Reserved.")
copyright_label.pack(side="bottom")

root.mainloop()