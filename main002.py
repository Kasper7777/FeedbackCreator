import os
from datetime import datetime
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Pt, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn  # Correct import for qn

# Configuration settings
config = {
    'percentage_mark_label': 'Percentage Mark:',
    'tutor_name': 'Dr. Kazber',  # Add the tutor's name
    'module_title': 'Introduction to Game Design',  # Add module title
    'module_code': 'GD101',  # Add module code
    'assignment_title': 'Assignment 1',  # Add assignment title
    'percent_of_module': '100%',  # Add percentage of module
    'font_sizes': {
        'title': 20,
        'year': 12,
        'comment': 14,
        'footer': 9
    }
}


# Helper function to find the first file with the given extension
def find_file(extension):
    for file in os.listdir('.'):
        if file.endswith(extension):
            return file
    return None


# Updated function to set cell border using proper namespaced attributes
def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for edge in ('top', 'left', 'bottom', 'right'):
        edge_data = kwargs.get(edge, None)
        if edge_data:
            edge_element = OxmlElement(f'w:{edge}')
            edge_element.set(qn('w:val'), 'single')
            edge_element.set(qn('w:sz'), str(edge_data * 8))  # width in eighths of a point
            edge_element.set(qn('w:space'), '0')
            edge_element.set(qn('w:color'), '000000')
            tcPr.append(edge_element)


# Helper function to add a centered heading
def add_centered_heading(doc, text, font_size=18):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt(font_size)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


# Helper function to extract the last table from a .docx document
def extract_last_table_from_docx(doc_path):
    doc = Document(doc_path)
    if doc.tables:
        return doc.tables[-1]
    return None


def append_table_to_document(target_doc, table):
    """Append the given table to the target document without modifications."""
    new_table = target_doc.add_table(rows=len(table.rows), cols=len(table.columns))
    new_table.autofit = False
    for row_idx, row in enumerate(table.rows):
        for cell_idx, cell in enumerate(row.cells):
            new_cell = new_table.cell(row_idx, cell_idx)
            new_cell.text = cell.text
            set_cell_border(new_cell, top=1, left=1, bottom=1, right=1)
            # Copying formatting
            for paragraph in cell.paragraphs:
                new_p = new_cell.paragraphs[0]
                for run in paragraph.runs:
                    new_run = new_p.add_run(run.text)
                    if run.bold:
                        new_run.bold = True
                    if run.italic:
                        new_run.italic = True
                    new_run.font.name = run.font.name
                    new_run.font.size = run.font.size
                new_p.alignment = paragraph.alignment


def create_feedback_form(student_name, student_mark, feedback, extra_table):
    doc = Document()

    # Set author to tutor's name
    core_properties = doc.core_properties
    core_properties.author = config['tutor_name']

    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Title
    add_centered_heading(doc, 'ASSIGNMENT FEEDBACK FORM', font_size=config['font_sizes']['title'])
    add_centered_heading(doc, '2024-25', font_size=config['font_sizes']['year'])

    # Add Student Info table
    table = doc.add_table(rows=3, cols=4)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in table.rows:
        for cell in row.cells:
            cell.width = Inches(1.5)

    # Set student information cells
    table.cell(0, 0).text = 'STUDENT:'
    table.cell(0, 1).text = student_name
    table.cell(0, 2).text = 'TUTOR:'
    table.cell(0, 3).text = config['tutor_name']
    table.cell(1, 0).text = 'MODULE TITLE:'
    table.cell(1, 1).text = config['module_title']
    table.cell(1, 2).text = 'MODULE CODE:'
    table.cell(1, 3).text = config['module_code']
    table.cell(2, 0).text = 'ASSIGNMENT:'
    table.cell(2, 1).text = config['assignment_title']
    table.cell(2, 2).text = '% of module:'
    table.cell(2, 3).text = config['percent_of_module']

    # Set table borders
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell, top=1, left=1, bottom=1, right=1)

    doc.add_paragraph()  # Add a break

    # Add "Overall Comment" box
    add_centered_heading(doc, 'OVERALL COMMENT', font_size=config['font_sizes']['comment'])
    comment_table = doc.add_table(rows=1, cols=1)
    comment_table.autofit = False
    comment_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    comment_table.cell(0, 0).width = Inches(6)
    comment_table.cell(0, 0).text = feedback if feedback else ' '
    set_cell_border(comment_table.cell(0, 0), top=1, left=1, bottom=1, right=1)
    doc.add_paragraph()  # Add a break

    # Add Marks and Date table
    table = doc.add_table(rows=1, cols=3)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for cell in table.rows[0].cells:
        cell.width = Inches(2)

    table.cell(0, 0).text = config['percentage_mark_label']
    mark_paragraph = table.cell(0, 1).paragraphs[0]
    mark_run = mark_paragraph.add_run(str(student_mark))
    mark_run.bold = True
    table.cell(0, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.cell(0, 2).text = 'Date:'
    table.cell(0, 2).paragraphs[0].add_run(datetime.now().strftime('%Y-%m-%d')).bold = True

    # Set table borders for Marks and Date table
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell, top=1, left=1, bottom=1, right=1)

    doc.add_paragraph()  # Add a break

    # Footer message
    footer = doc.add_paragraph()
    footer_run = footer.add_run(
        "NB All marks are provisional until confirmed by a formally constituted Board of Examiners")
    footer_run.font.size = Pt(config['font_sizes']['footer'])
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()  # Add a break

    # Add a page break after the footer message
    doc.add_page_break()

    # Append the extra table if it exists
    if extra_table:
        append_table_to_document(doc, extra_table)

    # Save the document
    file_name = f'Assignment_Feedback_Form_{student_name}.docx'
    doc.save(file_name)
    print(f"Saved: {file_name}")


# Find the Excel file in the current directory
excel_file = find_file('.xlsx')
if not excel_file:
    print("No Excel file found in the current directory.")
else:
    # Find the additional .docx file in the current directory
    docx_file = find_file('.docx')
    extra_table = extract_last_table_from_docx(docx_file) if docx_file else None

    # Read student names, marks, and feedback from Excel file
    df = pd.read_excel(excel_file)
    student_data = df.iloc[:, [0, 1, 2]]

    # Create a feedback form for each student
    for index, row in student_data.iterrows():
        student_name = row.iloc[0]
        student_mark = row.iloc[1]
        feedback = row.iloc[2]
        create_feedback_form(student_name, student_mark, feedback, extra_table)